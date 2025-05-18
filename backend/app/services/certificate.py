from docx import Document
import os
from sqlalchemy.orm import Session
from ...models.calibration import Calibration
from ...models.measurement import Measurement
from ...models.company import Company

def generate_certificate(calibration_id: int, db: Session):
    calibration = db.query(Calibration).filter(Calibration.id == calibration_id).first()
    measurements = db.query(Measurement).filter(Measurement.calibration_id == calibration_id).all()
    company = db.query(Company).filter(Company.id == calibration.company_id).first()
    
    # Placeholder: Assumes template.docx exists in certificates/
    # TODO: Provide actual template.docx to finalize this logic
    doc = Document()
    doc.add_heading("Calibration Certificate", 0)
    
    doc.add_paragraph(f"Serial Number: {calibration.serial_number}")
    doc.add_paragraph(f"Calibration Date: {calibration.calibration_date}")
    doc.add_paragraph(f"Company: {company.name if company else 'Unknown'}")
    
    table = doc.add_table(rows=len(measurements) + 1, cols=4)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "SSD (m)"
    hdr_cells[1].text = "Ref Dose Rate (mSv/h)"
    hdr_cells[2].text = "Corrected Dose"
    hdr_cells[3].text = "Calibration Factor"
    
    for i, measurement in enumerate(measurements, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = str(measurement.ssd)
        row_cells[1].text = str(measurement.ref_dose_rate_msv_h)
        row_cells[2].text = str(measurement.corrected_dose)
        row_cells[3].text = str(measurement.calibration_factor)
    
    doc.add_paragraph(f"Environmental Conditions:")
    doc.add_paragraph(f"Initial: Temp {calibration.initial_temperature}°C, Pressure {calibration.initial_pressure} kPa, Humidity {calibration.initial_humidity}%")
    doc.add_paragraph(f"Final: Temp {calibration.final_temperature}°C, Pressure {calibration.final_pressure} kPa, Humidity {calibration.final_humidity}%")
    
    output_path = f"certificates/certificate_{calibration_id}.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path