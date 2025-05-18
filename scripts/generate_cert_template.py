from docx import Document
import os

def generate_template():
    doc = Document()
    doc.add_heading("Calibration Certificate Template", 0)
    doc.add_paragraph("This is a placeholder template.")
    doc.add_paragraph("Fields to include:")
    doc.add_paragraph("- Serial Number\n- Calibration Date\n- Company\n- SSD Table\n- Environmental Conditions")
    
    output_path = "certificates/template.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Template saved to {output_path}")

if __name__ == "__main__":
    generate_template()